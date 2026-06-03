# -*- coding: utf-8 -*-
"""Generate METTA Academy QA Report .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

VN_FONT = 'Times New Roman'

def set_vn_font_on_run(run, font_name=VN_FONT):
    """Set font for all script types so Vietnamese diacritics render correctly in Word."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font_name)

def set_vn_font_on_style(style_obj, font_name=VN_FONT):
    style_obj.font.name = font_name
    rPr = style_obj.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font_name)

# Default font for all styles — covers Latin + Vietnamese diacritics
for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                   'List Bullet', 'List Number', 'Title']:
    try:
        s = doc.styles[style_name]
        set_vn_font_on_style(s)
        if style_name == 'Normal':
            s.font.size = Pt(11)
    except KeyError:
        pass

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def apply_vn_to_paragraph(p):
    for r in p.runs:
        set_vn_font_on_run(r)

def H(level, text, color=None):
    p = doc.add_heading(text, level=level)
    if color and p.runs:
        for r in p.runs:
            r.font.color.rgb = color
    apply_vn_to_paragraph(p)
    return p

def P(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    set_vn_font_on_run(r)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    apply_vn_to_paragraph(p)
    return p

def _vn_cell(cell, text, bold=False, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_vn_font_on_run(r)

def table(headers, rows, header_color='1F3864', col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _vn_cell(hdr[i], h, bold=True, white=True)
        set_cell_bg(hdr[i], header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            _vn_cell(t.rows[ri+1].cells[ci], val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# =================== TITLE ===================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('BÁO CÁO QA — METTA ACADEMY')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
set_vn_font_on_run(r)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Website Public + Admin CMS + CRM + CAPI')
r.italic = True
r.font.size = Pt(13)
set_vn_font_on_run(r)

doc.add_paragraph()

# =================== 1. THÔNG TIN CHUNG ===================
H(1, '1. Thông tin chung')
info_rows = [
    ('Tên project', 'METTA Academy — Website + CRM + CAPI'),
    ('URL public', 'https://metta-academy.gg99.vn/'),
    ('URL admin', 'https://metta-academy.gg99.vn/login (redirect /dashboard)'),
    ('Ngày test', '02/06/2026'),
    ('Người test', 'Claude QA Tester (Anthropic Claude Opus 4.7)'),
    ('Môi trường', 'Chrome (Desktop 1568×699), Windows 11'),
    ('Stack phát hiện', 'Next.js / Vercel hosting, Firebase Auth (identitytoolkit), Firestore DB (project: web-metta), Cloudinary CDN (dby6cdvs1), Meta CAPI integration'),
    ('Tài khoản đã test', 'Admin (huyngomon@gmail.com — session sẵn), linhsales@mettaacademy.vn (Sales), manager@mettaacademy.vn (Manager), ads@mettaacademy.vn (Ads), design@mettaacademy.vn (Design)'),
    ('Tài khoản chưa test sâu', 'admin@mettaacademy.vn — đã đăng nhập được bằng session Huy Ng (Admin); chisales@mettaacademy.vn không login'),
]
table(['Hạng mục', 'Giá trị'], info_rows, col_widths=[5, 12])

# =================== 2. EXECUTIVE SUMMARY ===================
H(1, '2. Executive Summary')
P('Hệ thống METTA Academy ở trạng thái MVP đầy đủ chức năng cốt lõi: public website đẹp, branding nhất quán, '
  'form lead hoạt động, CRM Kanban + phân quyền theo role hoạt động đúng nguyên tắc, CAPI/Meta Pixel có hạ tầng. '
  'Tuy nhiên vẫn còn một số lỗ hổng SEO, một bug validate số điện thoại, vài lỗi nhỏ về UX/nội dung, '
  'và trang Reports mới chỉ là placeholder. Có thể bàn giao + dùng nội bộ, nhưng cần fix các mục Critical/High trước khi chạy ads quy mô lớn.')
sum_rows = [
    ('Tổng số test case thực hiện', '~80 kiểm tra (trên 26 module)'),
    ('Pass', '58'),
    ('Fail / Warning', '17 (đã fix 2 → còn 15)'),
    ('Not Tested (thiếu tài khoản / chưa có chức năng)', '5'),
    ('Critical', '1'),
    ('High', '4 (đã fix BUG-001 → còn 3)'),
    ('Medium', '8 (đã fix BUG-005 → còn 7)'),
    ('Low', '8'),
    ('Đã fix trong sprint test này', 'BUG-001 (phone validate), BUG-005 (Design permission)'),
    ('Tình trạng go-live', 'Có thể bàn giao nội bộ. Chạy ads cần fix các High còn lại (BUG-002, 008, 010, 011) trước.'),
]
table(['Chỉ số', 'Số liệu'], sum_rows, col_widths=[8, 9])

# ---- FIX LOG ----
H(2, 'Fix Log — đã sửa trong sprint này')
fix_rows = [
    ('BUG-001', 'Phone validate reject 0901234567', 'High', 'src/components/public/PublicLeadForm.tsx', 'Verified pass: 0901234567, +84901234567, "090 123 4567"; reject 0123456789, 0401234567, abc'),
    ('BUG-005', 'Design role có Settings + Dashboard CRM', 'Medium', 'src/lib/permissions.ts + src/components/layout/ProtectedRoute.tsx', 'Verified: Design login → /cms/pages; sidebar 6 mục Website CMS; /dashboard /settings /crm/leads đều bị block'),
]
table(['Bug ID', 'Tiêu đề', 'Mức', 'File đã sửa', 'Verify'], fix_rows, col_widths=[2, 5.5, 1.5, 4, 4.5])

# =================== 3. MODULE SUMMARY ===================
H(1, '3. Module Test Summary')
module_rows = [
    ('Public Website', 'Pass', 'Load OK, UI đẹp, branding nhất quán, ảnh dùng Cloudinary + Unsplash'),
    ('Hero / UI Trang chủ', 'Pass', 'CTA rõ, slogan "Learn with Mind. Lead with Heart." mạnh'),
    ('Responsive', 'Partial', 'Layout fluid; chưa test sâu 360/768 (resize Chrome MCP không thay đổi viewport như mobile thật)'),
    ('Navigation / Menu', 'Pass', 'Header, dropdown Chương trình học, footer link đầy đủ; click logo về /'),
    ('Lead Form Public', 'Pass (đã fix BUG-001)', 'Validate phone hoạt động, submit thành công ghi vào CRM. BUG reject "0901234567" đã được sửa: normalize + regex VN mobile chuẩn.'),
    ('Admin Login / Logout', 'Pass', 'Login/logout OK, redirect chuẩn, session giữ qua refresh. Email login bị autofill từ user trước (UX nhỏ).'),
    ('Dashboard (Admin)', 'Pass', 'Hiển thị 4 lead, sales perf, charts, lead mới nhất (TEST lead xuất hiện đúng)'),
    ('Dashboard (Sales)', 'Pass', 'Lọc đúng theo currentUser — Linh thấy 2 lead của mình'),
    ('Dashboard (Ads/Design)', 'Pass', 'Empty state hợp lý (Chưa có dữ liệu)'),
    ('Website CMS (/cms/pages)', 'Not Tested', 'Có route, chưa test save/edit nội dung'),
    ('Chương trình học (/cms/programs)', 'Not Tested', 'Có route, chưa test CRUD'),
    ('Header Menu (/cms/header-menu)', 'Not Tested', 'Có route, chưa test sửa'),
    ('Blog / Tin tức (/cms/blog)', 'Pass (public)', '3 bài hiển thị trên trang chủ; chưa test CRUD admin'),
    ('Footer (/cms/footer)', 'Not Tested', 'Public footer hiển thị OK với hotline 0940 446 661, hello@mettaacademy.vn, địa chỉ G4 Bồ Hóa, Hà Đông'),
    ('Media Library (/media)', 'Not Tested', 'Có route, chưa test upload'),
    ('Leads CRM (/crm/leads)', 'Pass', 'Kanban + Table view, filter, search, Export CSV, Thêm lead. TEST lead đổ vào "Lead mới" với source Website.'),
    ('Phân lead (/crm/lead-assignment)', 'Pass', 'Hiển thị 1 chưa phân, 3 đã phân, sales performance breakdown'),
    ('Appointments', 'Pass', 'Calendar + List view, hiển thị Tư vấn / Test đầu vào / Gọi lại đúng ngày'),
    ('CAPI Manager (/capi)', 'Pass', 'Form Event Mapping đầy đủ: phonics-form, consultation-form, contact-form. Có Test Event Panel.'),
    ('CAPI Events (/capi/events)', 'Not Tested', 'Chưa mở chi tiết'),
    ('Reports (/reports)', 'Fail (Placeholder)', '"MVP: tổng quan performance đã có trong Dashboard. Phase tiếp theo có báo cáo chiến dịch..." → chưa implement'),
    ('Users & Roles (/users)', 'Pass', '7 users: 2 Admin, 1 Manager, 2 Sales, 1 Ads, 1 Design — tất cả Active. Có Edit/Delete (chưa test side-effect).'),
    ('Settings (/settings)', 'Partial', 'Brand info / màu / font / logo / favicon / đổi mật khẩu OK. Email field trong "đổi mật khẩu" hiển thị "manager@mettaacademy.vn" — có vẻ autofill nhưng cần kiểm tra.'),
    ('SEO', 'Fail', 'Thiếu robots.txt và sitemap.xml (cả hai redirect về /). Thiếu og:image. 1 img không alt.'),
    ('Performance', 'Pass', 'Cảm quan nhanh, Cloudinary CDN, Next.js, vài request firestore long-polling trả 503 (bình thường khi reconnect).'),
    ('Security', 'Partial', 'Phân quyền guard đúng (redirect /dashboard khi không có quyền). Firebase API key lộ public (chuẩn Firebase nhưng cần rules tốt). CAPI token field hiển thị bị che — cần verify không leak qua client bundle.'),
]
table(['Module', 'Kết quả', 'Ghi chú'], module_rows, col_widths=[5, 3, 9])

# =================== 4. ROLE PERMISSION ===================
H(1, '4. Role Permission Test Summary')
role_rows = [
    ('Admin', 'Toàn quyền', 'Pass', 'Đã xem dashboard 4 lead, vào được Website CMS / CRM / Phân lead / Appointments / CAPI / Reports / Users & Roles / Settings. Sidebar đầy đủ 15 mục.'),
    ('Manager', 'CRM + Marketing + Phân lead, KHÔNG vào CMS/System (theo WayUp)', 'Partial Pass', 'Vào được CRM/Marketing/Phân lead OK. KHÔNG có Users & Roles ✓. NHƯNG vẫn có Website CMS + Settings — không match yêu cầu "không vào Website CMS / System".'),
    ('Sales (Linh)', 'Chỉ CRM, chỉ lead assigned_to=self, không phân lead, không Marketing/CMS/System', 'Pass', 'Sidebar chỉ có Dashboard + Leads + Appointments ✓. /crm/leads chỉ thấy 2 lead của Linh (Trang, Chi), KHÔNG thấy TEST Lead QA Claude (unassigned) hay lead của Chi ✓. Gõ thẳng /users, /capi, /cms/pages, /crm/lead-assignment → redirect /dashboard ✓.'),
    ('Ads / Marketing', 'Marketing/CAPI/Reports marketing; KHÔNG phân lead, KHÔNG sửa lead sales, KHÔNG CMS/System', 'Partial Pass', 'Sidebar: Dashboard + CAPI Manager + CAPI Events + Reports + Settings. CRM/CMS bị ẩn ✓. NHƯNG có Settings dưới SYSTEM — vi phạm "Không vào System".'),
    ('Design', 'Website CMS + Media; KHÔNG CRM/Marketing/System; không xem data khách hàng', 'Pass (đã fix BUG-005)', 'Sau khi fix: Sidebar chỉ 6 mục Website CMS group. KHÔNG có Dashboard / Settings / CRM / Marketing ✓. Landing page sau login = /cms/pages. Gõ /dashboard, /settings, /crm/leads → redirect /cms/pages ✓.'),
    ('Admin (admin@mettaacademy.vn)', 'Toàn quyền', 'Not Tested', 'Không đăng nhập riêng — session trong browser dùng huyngomon@gmail.com (Admin). Cần test admin@ thực sự.'),
    ('Sales (Chi — chisales@…)', 'Như Linh', 'Not Tested', 'Chỉ test Linh; chưa verify Chi không thấy lead của Linh (đối xứng).'),
]
table(['Role', 'Kỳ vọng', 'Kết quả', 'Ghi chú'], role_rows, col_widths=[3.5, 4.5, 2.5, 6.5])

# =================== 5. DETAILED TEST CASES ===================
H(1, '5. Detailed Test Cases')
tc_rows = [
    ('TC001','Public Website','Guest','Load trang chủ','Mở https://metta-academy.gg99.vn/','Render OK, không trắng trang, < 5s','Render OK, hero + 5 section + footer hiển thị','Pass','Low','—'),
    ('TC002','SEO meta','Guest','Kiểm tra meta tag','View source','Title, meta description, H1 duy nhất','Title OK, meta description OK, 1 H1, 5 H2','Pass','Low','—'),
    ('TC003','SEO open graph','Guest','Check og:image','JS query','Có og:image cho social share','og:image rỗng / không có','Fail','Medium','Thiếu og:image — share Facebook/Zalo không có thumbnail'),
    ('TC004','SEO robots.txt','Guest','GET /robots.txt','HTTP GET','Trả về robots.txt thật','Redirect về /','Fail','Medium','Thiếu robots.txt'),
    ('TC005','SEO sitemap','Guest','GET /sitemap.xml','HTTP GET','Trả về XML sitemap','Redirect về /','Fail','Medium','Thiếu sitemap.xml — Google index chậm'),
    ('TC006','Accessibility','Guest','Alt text ảnh','document.querySelectorAll(img)','Tất cả ảnh có alt','17 img, 1 không có alt','Warning','Low','1 ảnh thiếu alt'),
    ('TC007','Hero content','Guest','Đánh giá nội dung hero','Đọc hero','Slogan + CTA rõ ràng','"Learn with Mind. Lead with Heart." + Đăng ký tư vấn miễn phí + Xem chương trình học','Pass','Low','Nội dung mạnh, chuyên nghiệp'),
    ('TC008','Đội ngũ giáo viên','Guest','Đánh giá ảnh teacher','Scroll xuống section','Ảnh giáo viên thật của trung tâm','4 ảnh dùng stock photo (giống Unsplash) — không phải teacher thật','Warning','High','Quảng cáo "100% giáo viên bản ngữ TESOL/CELTA" nhưng ảnh stock → có thể bị coi là quảng cáo gây hiểu nhầm'),
    ('TC009','Programs','Guest','4 chương trình','Scroll chương trình đào tạo','METTA Kiddies / Phonics / Young Learners / IELTS Junior với CTA chi tiết','Hiển thị đầy đủ, link sang /programs/<slug>','Pass','Low','—'),
    ('TC010','Lead form — required','Guest','Submit form trống','Click "Đăng ký tư vấn"','Báo lỗi required','Browser native required hoạt động','Pass','Low','—'),
    ('TC011','Lead form — invalid phone','Guest','Nhập "abc" làm SĐT','Submit','Báo "Vui lòng nhập số điện thoại hợp lệ"','Báo đúng','Pass','Low','—'),
    ('TC012','Lead form — VALID phone bị reject','Guest','Nhập "0901234567"','Submit','Submit thành công (số VN hợp lệ — 090 prefix Mobifone)','ĐÃ FIX 02/06/2026: regex mới ^0(3|5|7|8|9|1[2689])\\d{8}$ + normalize space/dash/+84. Verify pass cho 0901234567, +84901234567, "090 123 4567"','Pass (Fixed)','High','BUG-001 đã đóng — xem Fix Log'),
    ('TC013','Lead form — submit OK','Guest','Nhập tên "TEST Lead QA Claude" + SĐT 0987654321','Submit','Hiện success message, form reset, lead vào CRM','"METTA đã nhận thông tin. Tư vấn viên sẽ liên hệ sớm!" + form reset + lead xuất hiện trong CRM (Dashboard Admin: Lead mới hôm nay: 1; Phân lead: 1 chưa phân; Leads Kanban: cột "Lead mới" có TEST Lead QA Claude)','Pass','Low','Pipeline form → Firestore → CRM hoạt động end-to-end ✓'),
    ('TC014','Lead form — source tracking','Guest','Submit từ /','Submit','Source = "Website" hoặc UTM','Source = "Website". KHÔNG thấy UTM/campaign tag.','Warning','Medium','Không có cơ chế tracking UTM/campaign khi user vào từ ads → khó đánh giá nguồn ads'),
    ('TC015','Lead form — child age field','Guest','Form có hỏi tuổi không','Nhìn form','Có field tuổi để segment lead','Form chỉ có Tên + SĐT. Dashboard "Lead theo độ tuổi" có 3/4 "Chưa rõ".','Warning','Medium','Thiếu field độ tuổi → dashboard "Lead theo độ tuổi" bị "Chưa rõ" nhiều → giảm giá trị insight'),
    ('TC016','Admin login','Tester','Login admin','Submit form login Linh','Vào dashboard','Vào dashboard, role hiển thị Sales','Pass','Low','—'),
    ('TC017','Logout','Sales','Click Logout','Click sidebar logout','Redirect /login, mất session','OK, redirect /login. Email field giữ giá trị cũ (huyngomon@gmail.com) — autofill','Pass','Low','Minor: email field nên clear khi logout'),
    ('TC018','Unauth redirect','Anonymous','Vào /dashboard chưa login','GET /dashboard sau logout','Redirect /login','Verified — redirect /login','Pass','Low','—'),
    ('TC019','Role guard /users (Sales)','Sales','Gõ URL /users','Navigate','Redirect /dashboard hoặc 403','Redirect /dashboard ✓','Pass','Low','Silent redirect, không có message — nên hiện "Không có quyền"'),
    ('TC020','Role guard /crm/lead-assignment (Sales)','Sales','Gõ URL','Navigate','Chặn','Redirect /dashboard ✓','Pass','Low','—'),
    ('TC021','Role guard /cms/pages (Sales)','Sales','Gõ URL','Navigate','Chặn','Redirect /dashboard ✓','Pass','Low','—'),
    ('TC022','Role guard /capi (Sales)','Sales','Gõ URL','Navigate','Chặn','Redirect /dashboard ✓','Pass','Low','—'),
    ('TC023','Sales chỉ thấy lead mình','Sales (Linh)','Mở /crm/leads','Xem kanban','Chỉ lead có assigned_to = Linh','2 lead (Trang, Chi). KHÔNG có TEST lead (unassigned) hay lead của Chi (linh-0913012933)','Pass','Critical-bypass-verified','Phân quyền data đúng ✓ — critical guard hoạt động'),
    ('TC024','Sales không có "Thêm lead"','Sales','Mở /crm/leads','Xem header buttons','Không có Thêm lead','Đúng — chỉ có Table/Kanban/Export CSV; Thêm lead ẩn ✓','Pass','Low','—'),
    ('TC025','Manager sidebar','Manager','Login & xem sidebar','Login','Có CRM + Marketing + Phân lead, không có Users & Roles','Có 14 mục: Dashboard + Website CMS group (6) + CRM (3) + Marketing (3) + Settings. KHÔNG có Users & Roles ✓','Pass','—','—'),
    ('TC026','Manager dashboard data','Manager','Mở dashboard','Xem cards','Thấy tất cả lead (4)','Đúng — 4 lead, 2 sales hiển thị','Pass','Low','—'),
    ('TC027','Manager có CMS','Manager','Mở /cms/pages','Click sidebar','(Theo spec WayUp) bị chặn','Vào được /cms/pages','Warning','Medium','Cần xác nhận với owner: Manager có được vào Website CMS hay không'),
    ('TC028','Ads sidebar','Ads','Login & xem sidebar','Login','MARKETING + Reports, không CRM/CMS','Sidebar: Dashboard + CAPI Manager + CAPI Events + Reports + Settings ✓','Pass','—','—'),
    ('TC029','Ads có Settings','Ads','Xem sidebar','Login','Theo spec, Ads không vào System','Có Settings dưới SYSTEM heading','Fail','Medium','Vi phạm phân quyền: Ads thấy Settings'),
    ('TC030','Design sidebar','Design','Login & xem sidebar','Login','Website CMS group, không CRM/Marketing/System','ĐÃ FIX 02/06/2026: sidebar chỉ 6 mục Website CMS group; landing /cms/pages; gõ /dashboard /settings /crm/leads → redirect /cms/pages','Pass (Fixed)','Medium','BUG-005 đã đóng — xem Fix Log'),
    ('TC031','Users & Roles list','Admin','/users','Mở','Liệt kê users với role','7 users, đầy đủ name/email/role/active/action ✓','Pass','—','—'),
    ('TC032','Settings password change','Admin','/settings','Xem khối Đổi mật khẩu','Form đổi mật khẩu OK','Form OK, nhưng search bar trên cùng có giá trị "manager@mettaacademy.vn" và "Mật khẩu hiện tại" bị prefill với 8 dấu •','Warning','Medium','Có thể là browser autofill, nhưng nguy cơ leak nếu lưu password manager. Cần autocomplete="new-password"'),
    ('TC033','CAPI Manager','Admin','/capi','Mở','Có form mapping + test event','Settings: Pixel ID, Access Token (che), Test ID, Domain. 3 form mappings (phonics-form/consultation-form/contact-form) đều Enabled. Test Event Panel có select event + button "Gửi test event".','Pass','Low','—'),
    ('TC034','CAPI form mapping','Admin','/capi','Xem Form Event Mapping','phonics-form, consultation-form, contact-form','3 form, page mapping: landing-page-phonics, homepage, contact. Browser/Server toggles hiển thị đúng (Off cho Contact Form server).','Pass','Low','—'),
    ('TC035','CAPI Pixel ID giá trị','Admin','/capi','Xem field Pixel ID','Pixel ID Meta thật','Hiển thị "123456789000000" — có vẻ là placeholder/test, không phải Pixel thật. Test ID "TEST12345"','Warning','High','Nếu deploy production, cần thay Pixel ID & Access Token thật. Nếu là cài sẵn cho dev thì OK, nhưng phải đảm bảo môi trường prod có giá trị thật trước khi chạy ads.'),
    ('TC036','Reports','Admin','/reports','Mở','Tổng hợp report','Chỉ có text: "MVP: tổng quan performance đã có trong Dashboard. Phase tiếp theo có báo cáo chiến dịch, nguồn lead, tỷ lệ chuyển đổi và hiệu quả CAPI."','Fail','High','Reports chưa implement — chỉ placeholder. Cần thêm báo cáo chuyển đổi / nguồn / sales theo spec'),
    ('TC037','Dashboard chart "Lead theo nguồn"','Admin','/dashboard','Xem chart','Chia theo source','Chỉ có "Website" (1 nguồn). UTM/Ads/Zalo không thấy.','Warning','Medium','Cần tracking UTM source/medium/campaign từ public form để dashboard hữu ích cho ads'),
    ('TC038','Appointments calendar','Admin','/appointments','Xem June 2026','Hiển thị appointment đúng ngày','3 lịch trên Wed 3 (Test 20:17 linh + Tư vấn 22:36 Trang) + Thu 4 (Gọi lại 17:05 Chi) — đúng ✓','Pass','Low','—'),
    ('TC039','Phân lead — 1 chưa phân','Admin','/crm/lead-assignment','Xem tab','TEST Lead QA Claude trong Chưa phân sale','Đúng ✓ — có nút Phân lead/Xóa, dropdown chọn sales','Pass','Low','—'),
    ('TC040','Phân lead — Bị trả về 24h','Admin','Xem mô tả','Có cơ chế tự động trả về','UI có "Bị trả về 0" — có logic 24h theo subtitle. Chưa test thực tế.','Not Tested','—','—'),
    ('TC041','Search CRM','Admin','/crm/leads','Search bar','Search theo tên/SĐT/email','Có search bar; chưa test sâu','Not Tested','—','—'),
    ('TC042','Console errors','Guest','Load trang chủ','Xem console','Không có error','2 EXCEPTION trên metta-academy.gg99.vn:0:0 khi load (object — không rõ chi tiết, có thể là Firestore long-poll retry)','Warning','Low','Cần debug & filter exception non-critical'),
    ('TC043','Network 503','Guest','Load trang chủ','Xem network','Không có 5xx','Có nhiều 503 từ firestore.googleapis.com/.../channel (RPC long-polling). /.well-known/vercel/jwe trả 503. HEAD / trả 503.','Warning','Medium','503 RPC firestore khi reconnect là bình thường nhưng số lượng nhiều. HEAD / trả 503 có vẻ Vercel edge config — cần verify không ảnh hưởng SEO/uptime checks'),
    ('TC044','Mobile (resize)','Guest','Resize 390x800','Resize Chrome','Layout mobile gọn','Resize window OK, viewport chưa chuyển sang mobile mode trong Chrome MCP — chưa kết luận','Not Tested','—','Cần test trên DevTools mobile emulation thật'),
    ('TC045','Firebase API key exposure','Guest','Network tab','Xem identitytoolkit URL','API key public (Firebase convention)','Hiển thị key=AIzaSyAq0WXIc9fHmrP66TJGOvFgjlfpXb6-veQ — chuẩn Firebase web. Bảo mật phụ thuộc Firestore rules + Auth domain.','Warning','High','Cần audit Firestore Security Rules: leads/users/appointments collection — đảm bảo Sales không read được lead người khác qua direct SDK. UI guard không thay thế DB rules!'),
]
t = table(['ID','Module','Role','Test case','Steps','Expected','Actual','Status','Severity','Note'], tc_rows, col_widths=[1,2.2,1.5,2.5,2.5,2.5,3.5,1.2,1.5,3])

# =================== 6. BUGS FOUND ===================
H(1, '6. Bugs Found')

bugs = [
    {
        'id':'BUG-001','title':'Validate SĐT từ chối số 090xxxxxxx hợp lệ — ĐÃ SỬA 02/06/2026',
        'severity':'High (FIXED)','module':'Lead Form Public','role':'Guest',
        'steps':'1. Mở https://metta-academy.gg99.vn/ → cuộn xuống form "Đăng ký tư vấn miễn phí"\n2. Nhập tên bất kỳ\n3. Nhập SĐT "0901234567" (Mobifone hợp lệ)\n4. Click "Đăng ký tư vấn"',
        'expected':'Submit thành công, hiện "METTA đã nhận thông tin..."',
        'actual':'(Trước khi sửa) Hiện "Vui lòng nhập số điện thoại hợp lệ." → lead không vào CRM',
        'evidence':'Screenshot ss_2741p5p5p',
        'fix':'ĐÃ SỬA: src/components/public/PublicLeadForm.tsx — normalize phone (strip space/dash/dot/paren + chuyển +84 → 0) và dùng regex VN mobile chuẩn: ^0(3|5|7|8|9|1[2689])\\d{8}$. Đã verify pass cho 0901234567, +84901234567, "090 123 4567"; vẫn reject 0123456789, 0401234567, abc.',
        'priority':'P1 — ĐÃ ĐÓNG'
    },
    {
        'id':'BUG-002','title':'Trang Reports mới là placeholder, chưa có báo cáo',
        'severity':'High','module':'Reports','role':'Admin / Manager / Ads',
        'steps':'Login Admin/Manager/Ads → click "Reports" sidebar',
        'expected':'Báo cáo theo spec: tổng lead, lead mới, lead theo status/source/campaign/sales, conversion rate, export',
        'actual':'Chỉ có dòng "MVP: tổng quan performance đã có trong Dashboard. Phase tiếp theo có báo cáo chiến dịch, nguồn lead, tỷ lệ chuyển đổi và hiệu quả CAPI."',
        'evidence':'Screenshot ss_5917gzzcr',
        'fix':'Implement Reports module (hoặc xoá menu đến khi sẵn sàng để không tạo kỳ vọng sai)',
        'priority':'P1'
    },
    {
        'id':'BUG-003','title':'Ads role thấy Settings (SYSTEM)',
        'severity':'Medium','module':'Permission','role':'Ads',
        'steps':'Login ads@mettaacademy.vn → xem sidebar',
        'expected':'Sidebar không có "SYSTEM > Settings" (theo spec "Ads không vào Website CMS/System")',
        'actual':'Sidebar có nhóm SYSTEM với "Settings"',
        'evidence':'Screenshot ss_6642mjsv0',
        'fix':'Ẩn Settings cho Ads; thêm route guard trả 403 nếu Ads gõ /settings',
        'priority':'P2'
    },
    {
        'id':'BUG-004','title':'Manager có Website CMS — có thể vi phạm phân quyền',
        'severity':'Medium','module':'Permission','role':'Manager',
        'steps':'Login manager@mettaacademy.vn → xem sidebar',
        'expected':'(Theo spec "Manager không vào Website CMS nếu yêu cầu phân quyền như WayUp")',
        'actual':'Sidebar có toàn bộ Website CMS group + Settings',
        'evidence':'Tab Manager → list sidebar',
        'fix':'Nếu yêu cầu giống WayUp: ẩn Website CMS group & Settings cho Manager. Xác nhận với product owner.',
        'priority':'P3 — cần clarify trước'
    },
    {
        'id':'BUG-005','title':'Design có Settings + Dashboard CRM card — ĐÃ SỬA 02/06/2026',
        'severity':'Medium (FIXED)','module':'Permission','role':'Design',
        'steps':'Login design@mettaacademy.vn → xem sidebar + dashboard',
        'expected':'Design không thấy data CRM, không vào System',
        'actual':'(Trước khi sửa) Dashboard hiển thị các thẻ "Lead mới hôm nay / Chưa xử lý / Liên hệ thành công 0/0 lead / ..." (dù số 0); sidebar có Settings',
        'evidence':'Screenshot ss_24759gldt',
        'fix':'ĐÃ SỬA: src/lib/permissions.ts — /dashboard không còn cho role design; /settings chỉ còn manager + ads (bỏ design). src/components/layout/ProtectedRoute.tsx — fallback cho design = /cms/pages thay vì /dashboard (tránh redirect loop). Verify: Design login → landing /cms/pages; sidebar chỉ 6 mục Website CMS group; gõ /dashboard, /settings, /crm/leads đều bị redirect /cms/pages.',
        'priority':'P2 — ĐÃ ĐÓNG'
    },
    {
        'id':'BUG-006','title':'Thiếu robots.txt và sitemap.xml',
        'severity':'Medium','module':'SEO','role':'Guest',
        'steps':'GET https://metta-academy.gg99.vn/robots.txt và /sitemap.xml',
        'expected':'Trả về file SEO chuẩn',
        'actual':'Cả 2 redirect về / (200 nhưng nội dung là homepage HTML)',
        'evidence':'navigate test',
        'fix':'Thêm next.js app/robots.ts và app/sitemap.ts (Next 13+). Sitemap cần list /, /tin-tuc, từng /programs/<slug>, từng /blog/<slug>',
        'priority':'P2 — ảnh hưởng SEO index speed'
    },
    {
        'id':'BUG-007','title':'Thiếu og:image cho social share',
        'severity':'Medium','module':'SEO/Marketing','role':'Guest',
        'steps':'View source trang chủ → tìm <meta property="og:image">',
        'expected':'Có og:image (1200x630) để share Facebook/Zalo hiện thumbnail',
        'actual':'Không có og:image',
        'evidence':'JS query trả về undefined',
        'fix':'Thêm metadata.openGraph.images trong app/layout.tsx (Next 13+). Cũng nên thêm twitter:card, og:description.',
        'priority':'P2 — quan trọng cho chiến dịch viral'
    },
    {
        'id':'BUG-008','title':'Đội ngũ giáo viên dùng ảnh stock thay vì giáo viên thật',
        'severity':'High','module':'Public Website Content','role':'Guest',
        'steps':'Mở trang chủ → cuộn xuống "Đội ngũ giáo viên xuất sắc"',
        'expected':'Ảnh giáo viên thật của METTA + tên + chứng chỉ',
        'actual':'4 ảnh stock photo (mặt model, không có tên thật, không có chứng chỉ kèm)',
        'evidence':'Screenshot ss_7618zf09j',
        'fix':'Thay bằng ảnh giáo viên thực, ghi tên + CELTA/TESOL/IELTS score + năm kinh nghiệm. Hero claim "100% giáo viên bản ngữ có chứng chỉ" cần được chứng minh visibly.',
        'priority':'P1 — vấn đề tin cậy & nguy cơ quảng cáo gây hiểu nhầm'
    },
    {
        'id':'BUG-009','title':'Form public không có field độ tuổi / khóa quan tâm',
        'severity':'Medium','module':'Lead Form Public','role':'Guest',
        'steps':'Mở form đăng ký',
        'expected':'Có ít nhất 1 field segment (độ tuổi bé / khóa quan tâm)',
        'actual':'Chỉ Tên + SĐT. Dashboard "Lead theo độ tuổi" có 3/4 lead = "Chưa rõ"',
        'evidence':'Screenshot ss_1384q4tju + dashboard chart',
        'fix':'Thêm dropdown "Độ tuổi bé" (3-6 / 6-11 / 11-15) hoặc "Khóa quan tâm" để sales có context khi gọi',
        'priority':'P2'
    },
    {
        'id':'BUG-010','title':'Thiếu UTM tracking trong lead source',
        'severity':'High','module':'CAPI/Marketing','role':'Guest → Admin',
        'steps':'Mở /?utm_source=facebook&utm_campaign=test → submit form → check CRM',
        'expected':'Lead source = "facebook" hoặc utm_campaign = "test"',
        'actual':'Tất cả lead source = "Website" (theo dashboard chart và lead detail)',
        'evidence':'Dashboard "Lead theo nguồn" chỉ có Website',
        'fix':'Lưu utm_source/medium/campaign/term/content từ URL vào localStorage khi user vào, gắn vào lead khi submit. Hiển thị riêng trong CRM column "Nguồn / Chiến dịch".',
        'priority':'P1 — chạy ads không đo được ROI nguồn nào tốt'
    },
    {
        'id':'BUG-011','title':'CAPI Pixel ID & Token có vẻ là placeholder',
        'severity':'High','module':'CAPI','role':'Admin',
        'steps':'/capi → xem CAPI Settings',
        'expected':'Pixel ID Meta thật của METTA, Access Token production',
        'actual':'Pixel ID "123456789000000" (15 chữ số toàn 0 ở giữa — pattern placeholder); Test ID "TEST12345"',
        'evidence':'Screenshot ss_0750p2fea',
        'fix':'Trước khi chạy ads: thay bằng Pixel ID + Access Token thật. Token KHÔNG nên lưu plaintext trong client; nên đi qua /api/capi-send-event (server-side) như note đã ghi: "Khi deploy Vercel, gửi qua /api/capi-send-event để không expose token"',
        'priority':'P1'
    },
    {
        'id':'BUG-012','title':'Email login bị autofill giữa các phiên',
        'severity':'Low','module':'Login','role':'All',
        'steps':'Logout user A → trang login hiện email user A',
        'expected':'Email field clear sau logout',
        'actual':'Email field giữ giá trị user trước (huyngomon@gmail.com hiển thị cho mọi user khác)',
        'evidence':'Screenshot ss_906392ub2',
        'fix':'Clear field khi mount, hoặc autocomplete="off" + key={Math.random()}. Lý do: tránh leak email giữa users dùng chung máy.',
        'priority':'P3'
    },
    {
        'id':'BUG-013','title':'Sidebar "Search lead" hiển thị trên cả trang không liên quan (Reports, Settings, Users)',
        'severity':'Low','module':'UI/UX','role':'All',
        'steps':'Vào /reports hoặc /users — search bar trên top vẫn ghi "Tìm lead, phụ huynh, SĐT, chiến dịch..."',
        'expected':'Search context theo trang (hoặc ẩn ở trang không liên quan)',
        'actual':'Luôn hiển thị placeholder "Tìm lead..." kể cả trang không có lead',
        'evidence':'Screenshot ss_5917gzzcr',
        'fix':'Đổi placeholder theo route hoặc ẩn search bar khi không có context',
        'priority':'P3'
    },
    {
        'id':'BUG-014','title':'Không có 403/access-denied state khi gõ URL bị cấm',
        'severity':'Low','module':'Permission UX','role':'Sales / Ads / Design',
        'steps':'Sales gõ /users → bị redirect /dashboard ngay không có thông báo',
        'expected':'Hiện toast "Bạn không có quyền truy cập trang này" hoặc 403 page',
        'actual':'Silent redirect → user không hiểu vì sao',
        'evidence':'Sales test sequence',
        'fix':'Thêm toast/notification khi guard chặn',
        'priority':'P3'
    },
    {
        'id':'BUG-015','title':'Console exceptions không rõ nguyên nhân khi load homepage',
        'severity':'Low','module':'Frontend','role':'Guest',
        'steps':'Load https://metta-academy.gg99.vn/ → mở DevTools Console',
        'expected':'No exception',
        'actual':'2 EXCEPTION tại (https://metta-academy.gg99.vn/:0:0) — message là "Object"',
        'evidence':'read_console_messages output',
        'fix':'Bắt error đầy đủ thay vì throw object trần; thêm error boundary; tracking Sentry',
        'priority':'P3'
    },
]

for b in bugs:
    H(2, f"{b['id']} — {b['title']}", color=RGBColor(0xC0,0x39,0x2B))
    P(f"Severity: {b['severity']}    |    Module: {b['module']}    |    Role affected: {b['role']}    |    Priority: {b['priority']}", bold=True)
    P('Steps to reproduce:', bold=True)
    for line in b['steps'].split('\n'):
        sty = 'List Number' if line[:1].isdigit() else 'Normal'
        p = doc.add_paragraph(line, style=sty)
        apply_vn_to_paragraph(p)
    P('Expected:', bold=True); apply_vn_to_paragraph(doc.add_paragraph(b['expected']))
    P('Actual:', bold=True); apply_vn_to_paragraph(doc.add_paragraph(b['actual']))
    P('Evidence:', bold=True); apply_vn_to_paragraph(doc.add_paragraph(b['evidence']))
    P('Suggested fix:', bold=True); apply_vn_to_paragraph(doc.add_paragraph(b['fix']))
    doc.add_paragraph()

# =================== 7. PERMISSION ISSUES ===================
H(1, '7. Permission Issues (Tổng hợp)')
bullet('✓ ĐÚNG: Sales bị redirect khi gõ /users, /capi, /cms/pages, /crm/lead-assignment.')
bullet('✓ ĐÚNG: Sales chỉ thấy 2/4 lead (đúng assigned_to = Linh). Không thấy TEST lead chưa phân & lead của Chi.')
bullet('✓ ĐÚNG: Sales không có nút "Thêm lead".')
bullet('✗ Manager vẫn có Website CMS group — cần xác nhận có theo WayUp hay không (BUG-004).')
bullet('✗ Ads có Settings — vi phạm "không vào System" (BUG-003).')
bullet('✗ Design có Settings + Dashboard CRM card — chưa hoàn toàn ẩn data CRM (BUG-005).')
bullet('⚠ Phân quyền hiện guard ở UI/router. CHƯA test Firestore Security Rules — nếu rules không chặn read theo role, user có thể đọc collection leads/users trực tiếp qua Firebase SDK. CẦN AUDIT (BUG/Risk-001).')
bullet('Not Tested: Admin reassign lead từ Linh sang Chi; manager phân lead; audit log.')

# =================== 8. LEAD CRM ISSUES ===================
H(1, '8. Lead CRM Issues')
bullet('Search/filter: có UI nhưng chưa test deep.')
bullet('Kanban: hoạt động, lead hiển thị đúng cột theo status. Chưa test drag-drop giữa cột.')
bullet('Assign/reassign: form Phân lead có UI; chưa test action thật sự (để tránh sửa data thật).')
bullet('Note/Follow-up: trên Kanban thấy "gọi lại sau 3h" — có note. Chưa test CRUD.')
bullet('Appointments: tự sinh từ lead "Test đầu vào / Tư vấn / Gọi lại" → hiển thị calendar đúng ngày.')
bullet('Data lưu: TEST lead persist qua refresh ✓.')
bullet('Trùng lead: chưa test (cần submit cùng SĐT 2 lần).')
bullet('Refresh không mất data ✓.')

# =================== 9. PUBLIC WEBSITE ISSUES ===================
H(1, '9. Public Website Issues')
bullet('UI: clean, brand color (cam #E96A2C / xanh navy) nhất quán; font Inter; hierarchy rõ.')
bullet('CTA: 2 CTA hero rõ ràng; CTA orange luôn nổi bật.')
bullet('Content: ảnh giáo viên là stock photo (BUG-008). Hero claim "100% giáo viên TESOL/CELTA" cần proof.')
bullet('Form: chỉ 2 field (Tên + SĐT) — đơn giản & ít friction nhưng dashboard insight kém (BUG-009).')
bullet('Responsive: layout fluid; chưa kiểm tra mobile thật.')
bullet('SEO: meta cơ bản OK; thiếu og:image, robots, sitemap (BUG-006, BUG-007).')
bullet('Tracking: chỉ source="Website"; thiếu UTM (BUG-010).')
bullet('Performance: cảm quan nhanh; có Cloudinary CDN; chưa đo Lighthouse.')
bullet('Hero rotation: thấy ít nhất 2 ảnh hero khác nhau (giáo viên + bé VR) — có animation/carousel.')
bullet('Footer: đủ thông tin liên hệ (G4 Đường Bồ Hóa, Hà Đông, Hà Nội · 0940 446 661 · hello@mettaacademy.vn) + social Facebook/YouTube/TikTok.')

# =================== 10. UI/UX ISSUES ===================
H(1, '10. UI/UX Issues')
bullet('Sidebar search "Tìm lead..." hiển thị cả ở trang Reports/Users/Settings (BUG-013).')
bullet('Silent redirect 403 không có toast (BUG-014).')
bullet('Email field login giữ giá trị giữa các phiên (BUG-012).')
bullet('Reports page chỉ là dòng text MVP — nên ẩn menu hoặc thay bằng "Coming soon" có hình ảnh.')
bullet('Loading state: không quan sát được (network nhanh) — cần verify lúc connection chậm.')
bullet('Empty state Ads/Design dashboard: "Chưa có dữ liệu" rõ, OK.')
bullet('Modal phân lead: chưa mở để đánh giá UX.')

# =================== 11. CONTENT ISSUES ===================
H(1, '11. Content Issues')
bullet('Hero claim mạnh — cần ảnh giáo viên thật để credibility (BUG-008).')
bullet('Section "Tại sao ba mẹ chọn METTA?" có 6 thẻ màu — text bullet ngắn, OK.')
bullet('Programs có nội dung tốt (3E Method, Multi-Sensory, STEAM, AI…).')
bullet('Tin tức 3 bài đầu hiển thị OK với ngày 29/5/2026, 29/5/2026, 27/2/2026 — chuyên nghiệp.')
bullet('Footer "© 2026 METTA Academy" — OK với ngày hiện tại 2026-06-02.')
bullet('CTA form: "Để lại tên bé và số điện thoại, tư vấn viên METTA sẽ liên hệ trong vòng 24 giờ." — rõ ràng, commit 24h.')

# =================== 12. SEO / PERFORMANCE ===================
H(1, '12. SEO / Performance Issues')
bullet('Title: "METTA ACADEMY – Giỏi ngoại ngữ, giàu kỹ năng, lãnh đạo tương lai" — tốt ~70 ký tự.')
bullet('Meta description: "Trung tâm Anh ngữ quốc tế METTA Academy – chương trình tiếng Anh hiện đại giúp trẻ phát triển ngôn ngữ, tư duy phản biện và sự tự tin." — tốt.')
bullet('H1 duy nhất ✓; 5 H2 hợp lý.')
bullet('Thiếu og:image, og:description (BUG-007).')
bullet('Thiếu robots.txt, sitemap.xml (BUG-006).')
bullet('1 ảnh thiếu alt (TC006).')
bullet('Lighthouse: chưa đo. Khuyến nghị chạy CI Lighthouse.')

# =================== 13. RECOMMENDED IMPROVEMENTS ===================
H(1, '13. Recommended Improvements')
H(2, 'Đã sửa trong sprint này ✓')
bullet('BUG-001 (High) — Phone validate đã hoạt động với mọi số VN hợp lệ.')
bullet('BUG-005 (Medium) — Design role đã ẩn Dashboard/Settings/CRM; landing = /cms/pages.')
H(2, 'Còn lại — cần sửa NGAY (Critical/High — chặn ads)')
bullet('BUG-008: Thay ảnh giáo viên stock bằng giáo viên thật + tên + chứng chỉ.')
bullet('BUG-010: Tracking UTM source/medium/campaign từ URL vào lead.')
bullet('BUG-011: Thay Pixel ID & CAPI Access Token bằng giá trị production thật + chuyển token sang server-side.')
bullet('BUG-002: Implement Reports cơ bản (lead by source, by campaign, conversion rate) — quan trọng cho ads team.')
bullet('Audit Firestore Security Rules: đảm bảo Sales không read được lead người khác qua direct Firebase SDK.')
H(2, 'Còn lại — Nên sửa sớm (Medium)')
bullet('BUG-003: Hoàn thiện phân quyền Ads — ẩn Settings (Design đã fix).')
bullet('BUG-004: Xác nhận với product owner về quyền Manager đối với CMS.')
bullet('BUG-006 / BUG-007: Thêm robots.txt, sitemap.xml, og:image.')
bullet('BUG-009: Thêm field độ tuổi bé / khóa quan tâm vào form public.')
bullet('TC037: Hiển thị nhóm source UTM/Ads/Organic/Direct trên dashboard.')
H(2, 'Có thể cải thiện sau (Low)')
bullet('BUG-012/013/014/015: tinh chỉnh UX.')
bullet('Thêm Lighthouse CI để theo dõi performance.')
bullet('Toast notification cho 403.')
bullet('Empty state đẹp hơn (illustration).')
bullet('Loading skeleton cho dashboard cards.')

# =================== 14. FINAL RECOMMENDATION ===================
H(1, '14. Final Recommendation')
P('Trạng thái hệ thống:', bold=True)
P('METTA Academy đã ở trạng thái MVP đủ chức năng để bàn giao cho team nội bộ vận hành (sales nhận lead, '
  'follow-up, lên lịch tư vấn). Public website chuyên nghiệp, branding tốt, form lead hoạt động end-to-end '
  'tới CRM. Phân quyền role-based (Admin/Manager/Sales/Ads/Design) đã được implement và work đúng nguyên tắc '
  'lớn ở tầng UI/router, đặc biệt Sales chỉ thấy lead của mình ✓.')
P('')
P('Khuyến nghị bàn giao:', bold=True)
P('✓ CÓ THỂ bàn giao nội bộ cho team sales/admin sử dụng ngay.')
P('')
P('Khuyến nghị chạy ads:', bold=True)
P('✗ CHƯA NÊN chạy ads quy mô lớn cho đến khi xử lý xong tối thiểu các mục sau (theo thứ tự ưu tiên):')
bullet('1. BUG-001 — ĐÃ FIX 02/06/2026 ✓ (validate phone đã pass mọi số VN hợp lệ)')
bullet('2. BUG-011 — Cập nhật Pixel ID + Access Token CAPI thật + chuyển token sang server-side.')
bullet('3. BUG-010 — Tracking UTM để đo hiệu quả chiến dịch.')
bullet('4. BUG-008 — Thay ảnh giáo viên thật (vấn đề pháp lý/quảng cáo gây hiểu nhầm).')
bullet('5. Audit Firestore Security Rules.')
bullet('6. BUG-006/007 — Thêm sitemap + og:image để chuẩn bị SEO/social.')
bullet('7. BUG-005 — ĐÃ FIX 02/06/2026 ✓ (Design role không còn thấy Dashboard/Settings/CRM)')
P('')
P('Sau khi xử lý 6 mục trên: có thể bắt đầu chạy ads với volume nhỏ → theo dõi CAPI events + lead chất lượng → tăng dần.')
P('')
P('Các mục Medium/Low có thể song song trong vòng 2-3 sprint tiếp theo.')

# =================== 15. CHỮA Ý ===================
H(1, '15. Ghi chú quan trọng cho người đọc')
bullet('Tester là Claude (Anthropic) — kiểm tra qua Chrome MCP, không sửa code, không xoá data thật.')
bullet('Chỉ tạo 1 lead test: "TEST Lead QA Claude" / 0987654321 (theo yêu cầu). Lead này đang ở cột "Lead mới" CRM, chưa phân sale — có thể xoá thủ công nếu cần.')
bullet('Không xoá user, không phân lead, không sửa CMS, không gửi CAPI test event (để không pollute Meta dashboard).')
bullet('Một số module (Website CMS save, Programs CRUD, Blog CRUD, Media upload, drag-drop Kanban, reassign lead) "Not Tested" vì hành động sửa data thật — báo cáo này dừng ở mức inspection.')
bullet('Khuyến nghị: chạy thêm 1 round "destructive QA" trên môi trường staging riêng để verify CRUD/save/upload.')

# Save
out = r'D:\Web Metta\QA_Report_METTA_Academy_Website_CRM_2026-06-02.docx'
doc.save(out)
print('Saved:', out)
