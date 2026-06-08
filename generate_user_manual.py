# -*- coding: utf-8 -*-
"""Generate METTA Academy User Manual .docx with screenshots per role."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

IMG = Path(r'D:\Web Metta\qa\manual_screenshots')

doc = Document()
section = doc.sections[0]
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)

VN_FONT = 'Times New Roman'

def set_vn_run(run, font=VN_FONT):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for attr in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
        rFonts.set(qn(attr), font)

def set_vn_style(style, font=VN_FONT):
    style.font.name = font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for attr in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
        rFonts.set(qn(attr), font)

for n in ['Normal','Heading 1','Heading 2','Heading 3','Heading 4','List Bullet','List Number','Title']:
    try:
        s = doc.styles[n]
        set_vn_style(s)
        if n == 'Normal':
            s.font.size = Pt(10.5)
    except KeyError: pass

# Tighten spacing globally to reduce page count
def tighten(style_name, size=None, before=2, after=2, line=1.0):
    try:
        pf = doc.styles[style_name].paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line
        if size: doc.styles[style_name].font.size = Pt(size)
    except KeyError: pass

tighten('Normal', size=10.5, before=0, after=2, line=1.05)
tighten('List Bullet', before=0, after=1, line=1.02)
tighten('List Number', before=0, after=1, line=1.02)
tighten('Heading 1', size=16, before=8, after=4)
tighten('Heading 2', size=13, before=6, after=2)
tighten('Heading 3', size=11.5, before=4, after=2)
tighten('Heading 4', size=11, before=3, after=1)

def apply_p(p):
    for r in p.runs: set_vn_run(r)

def H(level, text, color=None):
    p = doc.add_heading(text, level=level)
    if color and p.runs:
        for r in p.runs: r.font.color.rgb = color
    apply_p(p); return p

def P(text='', bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        if size: r.font.size = Pt(size)
        set_vn_run(r)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet'); apply_p(p); return p

def step(n, text):
    # Bullet-style step (Word auto-numbering accumulates across the whole doc,
    # so we use bullets instead — clean and resets visually per section).
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); set_vn_run(r); return p

def img(name, width_cm=11.0, caption=None):
    path = IMG / f'{name}.png'
    if not path.exists():
        P(f'[Hình {name} chưa có]', italic=True); return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(2)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
    if caption:
        c = P(f'Hình: {caption}', italic=True, size=8.5)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(4)

def set_cell_bg(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc.append(shd)

def _vn_cell(cell, text, bold=False, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text)); r.bold = bold
    if white: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    set_vn_run(r)

def table(headers, rows, header_color='1F3864', col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        _vn_cell(t.rows[0].cells[i], h, bold=True, white=True)
        set_cell_bg(t.rows[0].cells[i], header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            _vn_cell(t.rows[ri+1].cells[ci], val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in t.rows: r.cells[i].width = Cm(w)
    return t

def check(text):
    p = doc.add_paragraph()
    r = p.add_run('☐  ')
    set_vn_run(r); r.font.size = Pt(12)
    r2 = p.add_run(text); set_vn_run(r2)
    return p

def page_break():
    doc.add_page_break()

def info_grid(items, fill='EEF2F7'):
    rows = (len(items) + 1) // 2
    t = doc.add_table(rows=rows, cols=2)
    for idx, (label, value) in enumerate(items):
        cell = t.rows[idx // 2].cells[idx % 2]
        cell.text = ''
        set_cell_bg(cell, fill)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(2)
        r = p.add_run(label.upper()); r.bold = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64); set_vn_run(r)
        p2 = cell.add_paragraph(); p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(value); r2.font.size = Pt(10); set_vn_run(r2)
    if len(items) % 2 == 1:
        last = t.rows[-1].cells[1]; last.text = ''; set_cell_bg(last, 'FFFFFF')
    return t

def callout(label, text, fill='FDEEE3'):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]; c.text = ''
    set_cell_bg(c, fill)
    p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + '  '); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); set_vn_run(r)
    r2 = p.add_run(text); r2.font.size = Pt(10); set_vn_run(r2)
    return t

# ============== COVER (compact) ==============
eyebrow = doc.add_paragraph()
r = eyebrow.add_run('METTA ACADEMY'); r.bold = True; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); set_vn_run(r)
eyebrow.paragraph_format.space_after = Pt(2)

ttl = doc.add_paragraph()
r = ttl.add_run('Hướng Dẫn Sử Dụng Hệ Thống Theo Vai Trò Người Dùng')
r.bold = True; r.font.size = Pt(24); set_vn_run(r)
ttl.paragraph_format.space_after = Pt(4)

sub = doc.add_paragraph()
r = sub.add_run('Tài liệu dành cho 5 nhóm người dùng: Admin, Manager, Sales, Ads/Marketing, Design. '
                'Nội dung tập trung vào thao tác sử dụng hệ thống, không đi sâu vào kỹ thuật/code.')
r.font.size = Pt(10.5); set_vn_run(r)
sub.paragraph_format.space_after = Pt(8)

info_grid([
    ('Website công khai', 'https://metta-academy.gg99.vn/'),
    ('Trang quản trị (đăng nhập)', 'https://metta-academy.gg99.vn/login'),
    ('Đối tượng đọc', 'Người vận hành trung tâm: Admin, Manager, Sales, Ads, Design'),
    ('Phiên bản tài liệu', 'v1.0 — cập nhật 02/06/2026'),
])
doc.add_paragraph().paragraph_format.space_after = Pt(2)
callout('Lưu ý:', 'Mọi tài khoản test dùng chung mật khẩu Metta123. Tài liệu mô tả hệ thống thực tế tại '
        'thời điểm 02/06/2026. Một số thao tác ghi/xóa dữ liệu chỉ nên thực hiện khi đã hiểu rõ hậu quả.')

H(2, 'Phân quyền theo vai trò')
table(['Chức năng','Admin','Manager','Sales','Ads','Design'], [
    ('Dashboard CRM','✓','✓','✓ (chỉ self)','✓','—'),
    ('Website CMS (Pages/Programs/Header/Blog/Footer)','✓','✓','—','—','✓'),
    ('Media Library','✓','✓','—','—','✓'),
    ('Leads CRM','✓','✓','✓ (chỉ assigned)','—','—'),
    ('Phân lead','✓','✓','—','—','—'),
    ('Appointments','✓','✓','✓','—','—'),
    ('CAPI Manager / Events','✓','✓','—','✓','—'),
    ('Reports','✓','✓','—','✓','—'),
    ('Users & Roles','✓','—','—','—','—'),
    ('Settings','✓','✓','—','✓','—'),
], col_widths=[6.5,1.7,1.7,1.7,1.7,1.7])

H(2, 'Mục lục')
toc = [
    '1. Đăng nhập & Tài khoản test',
    '2. Public Website (cho khách / phụ huynh)',
    '3. Hướng dẫn cho ADMIN',
    '4. Hướng dẫn cho MANAGER',
    '5. Hướng dẫn cho SALES',
    '6. Hướng dẫn cho ADS / MARKETING',
    '7. Hướng dẫn cho DESIGN',
    '8. Câu hỏi thường gặp & Xử lý lỗi',
]
for t in toc: bullet(t)
page_break()

# ============== PHẦN 1. ĐĂNG NHẬP ==============
H(1, 'PHẦN 1. Đăng nhập & Tài khoản test')

H(2, '1.1. Tài khoản test cho từng role')
P('Trang đăng nhập: https://metta-academy.gg99.vn/login — Tất cả tài khoản test dùng chung mật khẩu: Metta123', bold=True)

table(['Role','Email','Tên hiển thị','Phạm vi truy cập'], [
    ('Admin','admin@mettaacademy.vn','admin','Toàn quyền hệ thống'),
    ('Manager','manager@mettaacademy.vn','Manager','CRM + CMS + Marketing + Reports + Settings'),
    ('Sales (Linh)','linhsales@mettaacademy.vn','Linh','Chỉ Dashboard + Leads + Appointments của bản thân'),
    ('Sales (Chi)','chisales@mettaacademy.vn','Chi','Chỉ Dashboard + Leads + Appointments của bản thân'),
    ('Ads / Marketing','ads@mettaacademy.vn','ads','Dashboard + CAPI + Reports + Settings'),
    ('Design','design@mettaacademy.vn','design','Chỉ Website CMS group'),
], col_widths=[3.2, 5.5, 2.8, 5.5])

P('Lưu ý: Hai tài khoản Sales có quyền giống nhau nhưng mỗi người chỉ thấy lead được gán cho mình.', italic=True)

H(2, '1.2. Cách đăng nhập')
img('login_page', 11, 'Trang đăng nhập')
step(1, 'Truy cập https://metta-academy.gg99.vn/login')
step(2, 'Nhập Email tài khoản theo bảng trên')
step(3, 'Nhập Password: Metta123')
step(4, 'Nhấn nút "Đăng nhập" màu cam')
step(5, 'Hệ thống tự chuyển đến dashboard/CMS theo role của bạn')

H(3, 'Trường hợp quên mật khẩu')
step(1, 'Nhấn "Đổi mật khẩu / Quên mật khẩu" ngay dưới nút Đăng nhập')
step(2, 'Nhập email → hệ thống gửi link reset qua Firebase Auth')
step(3, 'Kiểm tra email và làm theo hướng dẫn')

H(3, 'Đăng xuất')
step(1, 'Nhìn xuống cuối sidebar bên trái')
step(2, 'Nhấn nút "Logout"')
step(3, 'Hệ thống về trang /login')

page_break()

# ============== PHẦN 3. PUBLIC WEBSITE ==============
H(1, 'PHẦN 2. Public Website (dành cho khách / phụ huynh)')

P('Phần này không cần đăng nhập. Là trang public của trung tâm METTA Academy.')

H(2, '2.1. Trang chủ')
img('public_home', 11, 'Trang chủ với hero rotation + CTA')
P('Trang chủ có các phần:')
bullet('Header: Logo + menu Giới thiệu, Chương trình học (dropdown), Đội ngũ giáo viên, Tin Tức, Liên hệ + nút "Đăng ký tư vấn"')
bullet('Hero: slogan "Learn with Mind. Lead with Heart." + 2 CTA: "Đăng ký tư vấn miễn phí" và "Xem chương trình học"')
bullet('Section Chương trình đào tạo: 4 chương trình (Kiddies, Phonics, Young Learners, IELTS Junior)')
bullet('Section "Tại sao ba mẹ chọn METTA?": 6 thẻ điểm khác biệt')
bullet('Section Đội ngũ giáo viên')
bullet('Section Tin tức & Sự kiện (3 bài mới nhất)')
bullet('Form Đăng ký tư vấn miễn phí')
bullet('Footer: thông tin liên hệ, mạng xã hội, link chính sách')

H(2, '2.2. Trang Chương trình học')
img('public_program_kiddies', 11, 'Trang chi tiết chương trình METTA Kiddies')
P('4 chương trình hiện có:')
bullet('METTA Kiddies (3-6 tuổi) — /programs/metta-kiddies')
bullet('METTA on Phonics (5-7 tuổi) — /programs/metta-on-phonics')
bullet('METTA Young Learners (7-12 tuổi) — /programs/metta-young-learner')
bullet('IELTS Junior (11-15 tuổi) — /programs/ielts-junior')

img('public_program_phonics', 11, 'METTA on Phonics')
img('public_program_young_learner', 11, 'METTA Young Learners')
img('public_program_ielts', 11, 'IELTS Junior')

H(2, '2.3. Trang Tin tức')
img('public_news_list', 11, 'Danh sách bài viết')
P('Click vào từng bài để xem chi tiết.')

H(2, '2.4. Form đăng ký tư vấn')
P('Form xuất hiện ở cuối trang chủ và bất kỳ trang chương trình nào.')
P('Hướng dẫn cho phụ huynh:', bold=True)
step(1, 'Nhập "Họ và tên bé" — bắt buộc')
step(2, 'Nhập "Số điện thoại" — bắt buộc, phải là số VN hợp lệ (10 chữ số bắt đầu 03/05/07/08/09 hoặc 02). Có thể nhập kèm dấu cách (ví dụ "090 123 4567") hoặc +84 (ví dụ "+84901234567"), hệ thống tự chuẩn hóa')
step(3, 'Nhấn nút "Đăng ký tư vấn" (cam)')
step(4, 'Hệ thống hiện "METTA đã nhận thông tin. Tư vấn viên sẽ liên hệ sớm!" và xóa form')

P('Sau khi submit:', bold=True)
bullet('Lead tự động vào CRM ở cột "Lead mới"')
bullet('Source được set là "Website"')
bullet('Admin/Manager nhận thông báo 1 lead chưa có PIC')
bullet('CAPI Pixel + Server gửi event "CompleteRegistration" về Meta')

page_break()

# ============== PHẦN 4. ADMIN ==============
H(1, 'PHẦN 3. Hướng dẫn cho ADMIN')
P('Tài khoản test: ', bold=True); doc.paragraphs[-1].add_run('admin@mettaacademy.vn / Metta123')
apply_p(doc.paragraphs[-1])

P('Admin có TOÀN QUYỀN với hệ thống — bao gồm cả Users & Roles.', italic=True)

H(2, 'Checklist tổng — Admin phải biết làm')
check('Đăng nhập / đăng xuất')
check('Xem dashboard tổng quan CRM')
check('Quản lý nội dung Website CMS (Pages, Programs, Header, Blog, Footer)')
check('Quản lý Media Library (upload, sử dụng ảnh)')
check('Xem & sửa toàn bộ lead trong CRM')
check('Phân lead cho sales')
check('Xem lịch hẹn của toàn bộ sales')
check('Cấu hình Meta Pixel + CAPI')
check('Xem báo cáo')
check('Quản lý Users & Roles (tạo/sửa/xóa user, đổi role)')
check('Cấu hình branding (logo, màu, font, đổi mật khẩu)')

H(2, '3.1. Dashboard')
img('admin_dashboard', 11, 'Dashboard Admin — tổng quan CRM')
P('Dashboard hiển thị:')
bullet('8 thẻ số liệu trên đỉnh: Lead mới hôm nay, Chưa xử lý, Quá hạn follow-up, Liên hệ thành công %, Test/Học thử %, Đã chuyển đổi, Mất lead')
bullet('Filter thời gian: Hôm nay / 7 ngày / 30 ngày / Tháng này / Tháng trước / Tùy chọn')
bullet('Filter Sales / Nguồn / Khóa học')
bullet('Banner "X lead chưa có PIC" (link sang Phân lead)')
bullet('Biểu đồ Trạng thái lead')
bullet('Bảng Hiệu suất theo Sales (PIC): NHẬN, ĐÃ GỌI, % LH, TEST/HT, CHỐT, % CHỐT, MẤT, BỊ TRẢ VỀ')
bullet('Chart Xu hướng Lead theo ngày + Lead theo nguồn')
bullet('Card Lịch test/tư vấn sắp tới + Việc cần làm hôm nay + Lead mới nhất')

H(2, '3.2. Website CMS')

H(3, '3.2.1. Pages (Trang)')
img('admin_cms_pages', 11, 'Danh sách pages CMS')
P('Hướng dẫn quản lý trang:')
step(1, 'Sidebar → "Website CMS"')
step(2, 'Xem danh sách trang (Title, Slug, Status: draft/published, Updated)')
step(3, 'Click nút "Tạo page mới" để tạo trang mới')
step(4, 'Click vào title của trang để vào trình soạn thảo (Page Editor)')
step(5, 'Trong Page Editor: chỉnh từng section, lưu draft hoặc Publish')
step(6, 'Filter Draft / Published ở đầu danh sách')
step(7, 'Nút "Khôi phục dữ liệu mẫu CMS" để reset về dữ liệu seed (cẩn thận!)')

H(3, '3.2.2. Chương trình học')
img('admin_cms_programs', 11, 'Chương trình học CMS')
step(1, 'Sidebar → "Chương trình học"')
step(2, 'Xem 4 chương trình hiện có')
step(3, 'Sửa từng chương trình: tên, mô tả, ảnh đại diện, lợi ích, lịch học')
step(4, 'Bấm Save sau mỗi lần sửa')

H(3, '3.2.3. Header Menu')
img('admin_cms_header_menu', 11, 'Quản lý menu header')
step(1, 'Sidebar → "Header Menu"')
step(2, 'Sửa label hoặc link của từng mục menu')
step(3, 'Kéo thả để đổi thứ tự')
step(4, 'Bật/tắt mục menu')
step(5, 'Lưu thay đổi')

H(3, '3.2.4. Blog / Tin tức')
img('admin_cms_blog', 11, 'Blog admin')
step(1, 'Sidebar → "Blog / Tin tức"')
step(2, 'Xem danh sách bài viết')
step(3, 'Tạo bài mới: tiêu đề, slug, nội dung (rich text), ảnh đại diện, trạng thái draft/published')
step(4, 'Sửa bài hiện có')
step(5, 'Bài published xuất hiện ở /tin-tuc và trang chủ')

H(3, '3.2.5. Footer')
img('admin_cms_footer', 11, 'Footer CMS')
step(1, 'Sidebar → "Footer"')
step(2, 'Sửa hotline, email, địa chỉ trung tâm')
step(3, 'Sửa link mạng xã hội (Facebook, YouTube, TikTok)')
step(4, 'Sửa nội dung các cột link footer')
step(5, 'Lưu — public website cập nhật ngay')

H(3, '3.2.6. Media Library')
img('admin_media', 11, 'Media Library')
step(1, 'Sidebar → "Media Library"')
step(2, 'Upload ảnh mới (lưu vào Cloudinary)')
step(3, 'Copy URL ảnh để dùng trong CMS')
step(4, 'Xóa ảnh không dùng')
step(5, 'Lưu ý: file lớn (> 5MB) sẽ tải lâu — nén ảnh trước khi upload')

H(2, '3.3. CRM')

H(3, '3.3.1. Leads CRM')
img('admin_leads', 11, 'Leads CRM Kanban view')
P('Tính năng:')
bullet('2 view: Table / Kanban (mặc định Kanban)')
bullet('Filter: status, source, khóa, sales, khoảng ngày tạo')
bullet('Search bar: tìm theo tên/SĐT/email')
bullet('Export CSV: xuất danh sách hiện tại ra file CSV')
bullet('Thêm lead: tạo lead thủ công (khi nhận từ điện thoại/Zalo)')
bullet('Kéo thả lead giữa các cột status (trong Kanban)')
bullet('Click vào lead để mở chi tiết')

P('Cột Kanban:')
bullet('Lead mới — vừa từ form public, chưa được liên hệ')
bullet('Đã liên hệ — sales đã gọi & nhận')
bullet('Chưa nghe máy — sales gọi không bắt máy')
bullet('Đã hẹn tư vấn — đặt lịch tư vấn')
bullet('Đã tư vấn/Đặt lịch test — đã tư vấn xong, đặt lịch test đầu vào')
bullet('Đã test/Học thử — đã hoàn thành test')
bullet('Đã đăng ký học — chuyển đổi thành học viên')
bullet('Mất lead — không chốt được')

H(3, '3.3.2. Phân lead')
img('admin_lead_assignment', 11, 'Phân lead — màn hình admin')
P('Mục đích: gán lead chưa có PIC cho sales cụ thể.')
step(1, 'Sidebar → "Phân lead"')
step(2, 'Xem 3 tab: Chưa phân sale / Bị trả về / Đã phân sale')
step(3, 'Chọn lead chưa phân (tick checkbox)')
step(4, 'Chọn sales từ dropdown "Chọn sales nhận lead"')
step(5, 'Nhấn nút "Phân lead" (cam)')
step(6, 'Lead chuyển sang tab "Đã phân sale" và sales được thông báo')

P('Cơ chế tự trả về sau 24h:', bold=True)
bullet('Nếu sales được phân lead nhưng KHÔNG cập nhật status trong 24h, lead tự động về tab "Bị trả về"')
bullet('Admin/Manager có thể reassign cho sales khác')

P('Hiệu suất theo Sales (bảng dưới):')
bullet('SALES — tên sales')
bullet('ĐANG NHẬN — số lead đang cầm')
bullet('ĐÃ LIÊN HỆ — đã liên hệ xong')
bullet('CHỐT — đã chuyển đổi')
bullet('MẤT — lead mất')
bullet('BỊ TRẢ VỀ — số lần bị trả về')

H(3, '3.3.3. Appointments (Lịch hẹn)')
img('admin_appointments', 11, 'Calendar lịch hẹn')
P('Tính năng:')
bullet('2 view: Calendar (mặc định) / List')
bullet('Filter: khoảng ngày, tất cả sales, loại lịch (Test đầu vào / Tư vấn / Gọi lại)')
bullet('Click vào ô ngày để thêm lịch hẹn')
bullet('Click vào lịch sẵn có để sửa/xóa')
bullet('Admin/Manager thấy lịch của toàn bộ sales')

H(2, '3.4. CAPI Manager')
img('admin_capi', 11, 'CAPI Settings')

P('Tính năng:')
bullet('CAPI Settings: Pixel ID, Access Token, Test Event Code, Domain')
bullet('Toggle: enableBrowserPixel, enableServerCapi, enableDeduplication')
bullet('Test Event Panel: gửi event test (Lead/CompleteRegistration/...) để verify')
bullet('Form Event Mapping: bảng map form ID → event Meta + trạng thái Browser/Server')
bullet('Recent CAPI Logs: lịch sử event gần đây')

P('Hướng dẫn cấu hình lần đầu:', bold=True)
step(1, 'Đăng nhập Meta Business → Events Manager → lấy Pixel ID')
step(2, 'Tạo Access Token trong Meta CAPI Settings')
step(3, 'Vào /capi → dán Pixel ID và Access Token')
step(4, 'Bật enableServerCapi để token gửi từ server (an toàn hơn)')
step(5, 'Nhấn "Lưu CAPI"')
step(6, 'Nhấn "Gửi test event" → kiểm tra Meta Events Manager để xác nhận đã nhận')

img('admin_capi_events', 11, 'CAPI Events log')
P('CAPI Events: xem log từng event đã gửi (tên event, status, timestamp, payload).')

H(2, '3.5. Reports')
img('admin_reports', 11, 'Reports (MVP — chưa hoàn thiện)')
P('Trang Reports hiện ở trạng thái MVP placeholder. Báo cáo chi tiết (chiến dịch, nguồn lead, tỷ lệ chuyển đổi, CAPI) sẽ có ở phase tiếp theo. Tạm thời dùng Dashboard.', italic=True)

H(2, '3.6. Users & Roles (chỉ Admin)')
img('admin_users', 11, 'Quản lý users')
P('Tính năng:')
bullet('Danh sách users: name, email, role, active status')
bullet('Tạo user mới (nút "Thêm user")')
bullet('Sửa user (đổi tên, role, active)')
bullet('Xóa user (nút thùng rác)')
bullet('Role available: admin, manager, sales, ads, design')

P('Hướng dẫn tạo user mới:', bold=True)
step(1, 'Sidebar → "Users & Roles"')
step(2, 'Nhấn nút "Thêm user" (cam)')
step(3, 'Nhập email, mật khẩu tạm, tên hiển thị')
step(4, 'Chọn role phù hợp')
step(5, 'Bấm Save → user được tạo trong Firebase Auth + Firestore')
step(6, 'Gửi email và mật khẩu cho người mới — yêu cầu họ đổi mật khẩu lần đầu')

P('Lưu ý quan trọng:', bold=True)
bullet('Không xóa Admin cuối cùng — sẽ mất quyền truy cập')
bullet('Đổi role chỉ áp dụng sau khi user logout/login lại')
bullet('Vô hiệu hóa (active=false) thay vì xóa nếu nghỉ tạm')

H(2, '3.7. Settings')
img('admin_settings', 11, 'Settings — branding & password')
P('Tính năng:')
bullet('Đổi mật khẩu cá nhân (mật khẩu hiện tại + mật khẩu mới)')
bullet('Brand Name, Logo URL, Favicon URL, Font Family')
bullet('Primary Color / Secondary Color / Accent Color (color picker)')

P('Hướng dẫn:', bold=True)
step(1, 'Sidebar → "Settings"')
step(2, 'Đổi mật khẩu: nhập password hiện tại, password mới, nhập lại → Lưu')
step(3, 'Đổi logo: upload ảnh 240×104px nền trong suốt → copy URL từ Media Library → dán vào "Logo URL"')
step(4, 'Đổi favicon: ảnh 64×64px PNG hoặc ICO')
step(5, 'Đổi màu sắc: dùng color picker')
step(6, 'Bấm Save — thay đổi áp dụng ngay trên public website')

page_break()

# ============== PHẦN 5. MANAGER ==============
H(1, 'PHẦN 4. Hướng dẫn cho MANAGER')
P('Tài khoản test: ', bold=True); doc.paragraphs[-1].add_run('manager@mettaacademy.vn / Metta123')
apply_p(doc.paragraphs[-1])

P('Manager có quyền tương tự Admin nhưng KHÔNG được quản lý Users & Roles.', italic=True)

H(2, 'Checklist tổng — Manager phải biết làm')
check('Đăng nhập / đăng xuất')
check('Xem dashboard tổng quan CRM (đầy đủ data như Admin)')
check('Quản lý nội dung Website CMS')
check('Xem & sửa toàn bộ lead trong CRM')
check('Phân lead cho sales (chức năng cốt lõi của Manager)')
check('Quản lý lịch hẹn của sales')
check('Cấu hình CAPI Manager')
check('Xem Reports')
check('Đổi mật khẩu cá nhân trong Settings')

H(2, '4.1. Dashboard')
img('manager_dashboard', 11, 'Dashboard Manager — toàn bộ data như Admin')
P('Manager thấy tất cả số liệu như Admin: total lead, sales performance, charts.')

H(2, '4.2. Leads CRM (đầy đủ)')
img('manager_leads', 11, 'Leads CRM — manager view')
P('Manager có thể:', bold=True)
bullet('Xem toàn bộ lead (không bị giới hạn theo assigned_to)')
bullet('Sửa thông tin lead, đổi status, thêm note')
bullet('Tạo lead thủ công')
bullet('Export CSV')

H(2, '4.3. Phân lead — chức năng chính của Manager')
img('manager_lead_assignment', 11, 'Phân lead — manager view')
P('Quy trình hằng ngày của Manager:', bold=True)
step(1, 'Mỗi sáng vào "Phân lead" → kiểm tra tab "Chưa phân sale"')
step(2, 'Phân đều lead cho các sales theo nguyên tắc round-robin hoặc theo khóa quan tâm')
step(3, 'Kiểm tra tab "Bị trả về" → reassign các lead bị quá hạn 24h')
step(4, 'Theo dõi bảng "Hiệu suất theo Sales" để biết sales nào đang quá tải/đang free')

H(2, '4.4. Appointments')
img('manager_appointments', 11, 'Calendar manager view')
P('Manager thấy lịch của toàn bộ sales. Hữu ích để:')
bullet('Theo dõi sales nào có nhiều lịch tư vấn trong ngày')
bullet('Cảnh báo khi sales bỏ lịch')
bullet('Sắp xếp lại lịch khi cần')

H(2, '4.5. CAPI Manager')
img('manager_capi', 11, 'CAPI Manager — manager view')
P('Manager có quyền vào CAPI để phối hợp với Ads team. Không cần cấu hình từ đầu — thường chỉ vào để debug khi event không bắn.')

H(2, '4.6. Reports')
img('manager_reports', 11, 'Reports')

page_break()

# ============== PHẦN 6. SALES ==============
H(1, 'PHẦN 5. Hướng dẫn cho SALES')
P('Tài khoản test: ', bold=True); doc.paragraphs[-1].add_run('linhsales@mettaacademy.vn / Metta123 (hoặc chisales@…)')
apply_p(doc.paragraphs[-1])

P('Sales chỉ thấy lead được gán cho mình. Không thấy lead của sales khác. Không phân lead. Không vào CMS/Marketing/Users.', italic=True)

H(2, 'Checklist tổng — Sales phải biết làm')
check('Đăng nhập / đăng xuất')
check('Xem dashboard cá nhân (chỉ lead của mình)')
check('Mở danh sách lead được gán')
check('Cập nhật status lead (Lead mới → Đã liên hệ → Đã hẹn tư vấn → Đã test → Đã đăng ký)')
check('Thêm note cho lead sau mỗi lần gọi')
check('Tạo lịch hẹn (Tư vấn / Test đầu vào / Gọi lại)')
check('Theo dõi lịch hẹn trong ngày')
check('Phản hồi NHANH các lead bị "Bị trả về" để không mất lead')

H(2, '5.1. Dashboard cá nhân')
img('sales_dashboard', 11, 'Dashboard Sales — chỉ lead của mình')
P('Dashboard hiển thị các con số CHỈ TÍNH TRÊN LEAD CỦA SALES:')
bullet('Tổng lead nhận, % liên hệ thành công, % chốt, đã chuyển đổi, mất lead')
bullet('Bảng "Hiệu suất theo Sales" chỉ hiển thị bản thân')
bullet('Lead mới nhất — danh sách lead vừa được gán')

H(2, '5.2. Leads CRM')
img('sales_leads', 11, 'Leads Kanban — Sales chỉ thấy lead của mình')
P('Sales chỉ thấy lead có assigned_to = mình. Không thấy lead chưa phân hoặc lead của sales khác.', bold=True)

P('Quy trình xử lý 1 lead chuẩn:', bold=True)
step(1, 'Vào Leads CRM → tìm lead mới trong cột "Lead mới"')
step(2, 'Click vào lead để mở chi tiết')
step(3, 'Gọi điện cho phụ huynh theo SĐT')
step(4, 'Nếu nghe máy: kéo thẻ sang cột "Đã liên hệ" + thêm note nội dung cuộc gọi')
step(5, 'Nếu không nghe máy: kéo sang "Chưa nghe máy" + tạo lịch "Gọi lại" sau 2-3 giờ')
step(6, 'Nếu phụ huynh muốn tư vấn sâu: kéo sang "Đã hẹn tư vấn" + tạo Appointment "Tư vấn"')
step(7, 'Sau tư vấn: kéo sang "Đã tư vấn/Đặt lịch test" + tạo Appointment "Test đầu vào"')
step(8, 'Sau test: kéo sang "Đã test/Học thử" + nhập kết quả')
step(9, 'Nếu chốt: kéo sang "Đã đăng ký học"')
step(10, 'Nếu không chốt: kéo sang "Mất lead" + ghi lý do trong note')

P('⚠ LƯU Ý 24H:', bold=True)
bullet('Lead được phân nhưng KHÔNG cập nhật status trong 24h sẽ tự động trả về Manager.')
bullet('Để giữ lead: chỉ cần cập nhật status (kéo qua cột mới) hoặc thêm note.')

H(2, '5.3. Appointments')
img('sales_appointments', 11, 'Lịch hẹn cá nhân')
P('Sales thấy lịch hẹn liên quan đến lead của mình.')
P('Quy trình hàng ngày:', bold=True)
step(1, 'Đầu ngày: vào Appointments để xem lịch trong ngày')
step(2, 'Chuẩn bị tài liệu cho từng lịch tư vấn/test')
step(3, 'Sau khi hoàn thành lịch: cập nhật status lead tương ứng')
step(4, 'Nếu phụ huynh dời lịch: edit appointment đổi ngày giờ')

H(2, '5.4. Best practices cho Sales')
bullet('Gọi lead mới trong vòng 30 phút từ khi nhận để tăng tỷ lệ chốt')
bullet('Luôn thêm note sau mỗi cuộc gọi để Manager/đồng nghiệp hiểu context khi cần hỗ trợ')
bullet('Cập nhật status NGAY sau cuộc gọi — không để cuối ngày làm gộp')
bullet('Theo dõi % liên hệ thành công trên Dashboard để self-improve')

page_break()

# ============== PHẦN 7. ADS ==============
H(1, 'PHẦN 6. Hướng dẫn cho ADS / MARKETING')
P('Tài khoản test: ', bold=True); doc.paragraphs[-1].add_run('ads@mettaacademy.vn / Metta123')
apply_p(doc.paragraphs[-1])

P('Ads/Marketing tập trung vào: cấu hình Meta Pixel + CAPI, theo dõi event log, xem báo cáo marketing. Không vào CMS, không xem chi tiết khách hàng.', italic=True)

H(2, 'Checklist tổng — Ads phải biết làm')
check('Đăng nhập / đăng xuất')
check('Xem dashboard tổng quan (không thấy chi tiết khách hàng)')
check('Cấu hình Meta Pixel ID và CAPI Access Token')
check('Tạo và kiểm tra Form Event Mapping')
check('Gửi test event để verify tracking')
check('Xem log CAPI events')
check('Xem báo cáo marketing')
check('Đổi mật khẩu cá nhân')

H(2, '6.1. Dashboard')
img('ads_dashboard', 11, 'Dashboard Ads — empty state nếu chưa có data marketing')
P('Dashboard cho Ads role chỉ hiển thị metrics tổng (không drill-down vào tên/SĐT lead).')

H(2, '6.2. CAPI Manager — chức năng cốt lõi của Ads')
img('ads_capi', 11, 'CAPI Manager')

P('Phần 1: CAPI Settings', bold=True)
table(['Field','Mô tả'], [
    ('Pixel ID','Số ID Meta Pixel (lấy từ Events Manager)'),
    ('Access Token','Token CAPI (lấy từ Events Manager > Conversions API)'),
    ('Test Event Code','Mã test để Meta nhận diện event test (TEST12345)'),
    ('Domain','Domain website (https://mettaacademy.vn)'),
    ('enableBrowserPixel','Bật pixel browser-side (mặc định bật)'),
    ('enableServerCapi','Bật CAPI server-side (an toàn token)'),
    ('enableDeduplication','Bật dedupe để Meta không double-count'),
], col_widths=[5, 12])

P('Phần 2: Test Event Panel', bold=True)
step(1, 'Chọn event muốn test (Lead, CompleteRegistration, Contact, ViewContent…)')
step(2, 'Nhấn "Gửi test event"')
step(3, 'Mở tab Meta Events Manager → Test Events trong vài giây')
step(4, 'Verify event xuất hiện với Test Event Code khớp')

P('Phần 3: Form Event Mapping', bold=True)
P('Bảng map từng form trên website → event Meta. Hiện đã có 3 form mapping mặc định:')
table(['Form','Event Meta','Page','Browser','Server'], [
    ('phonics-form','Lead','/landing-page-phonics','On','On'),
    ('consultation-form','CompleteRegistration','/ (homepage)','On','On'),
    ('contact-form','Contact','/contact','On','Off'),
], col_widths=[4.5, 4.5, 3.5, 1.5, 1.5])

P('Thêm/sửa form mapping khi thêm form mới trên website.')

H(2, '6.3. CAPI Events log')
img('ads_capi_events', 11, 'Log từng event đã bắn')
P('Dùng để debug khi event không về Meta:')
bullet('Tên event, status (success/fail), timestamp')
bullet('Payload gửi đi')
bullet('Response Meta trả về')
bullet('Filter theo ngày/event/status')

H(2, '6.4. Reports')
img('ads_reports', 11, 'Reports — Ads view')
P('Hiện chỉ là placeholder. Phase tiếp theo sẽ có: ROAS, CPL, conversion rate theo campaign.')

H(2, '6.5. Best practices cho Ads')
bullet('Trước khi launch chiến dịch: chạy ít nhất 5 test event qua Test Event Panel để xác nhận pipeline hoạt động')
bullet('Theo dõi CAPI Events log hàng ngày — fail event > 5% là dấu hiệu cần debug')
bullet('Mỗi campaign mới nên có UTM riêng (utm_campaign, utm_source, utm_medium) để track ROI')
bullet('Không lưu Access Token vào client bundle — luôn dùng enableServerCapi=true')

page_break()

# ============== PHẦN 8. DESIGN ==============
H(1, 'PHẦN 7. Hướng dẫn cho DESIGN')
P('Tài khoản test: ', bold=True); doc.paragraphs[-1].add_run('design@mettaacademy.vn / Metta123')
apply_p(doc.paragraphs[-1])

P('Design role chỉ truy cập Website CMS group. Không vào Dashboard CRM, không xem data khách hàng, không vào Marketing/System.', italic=True)
P('Sau khi đăng nhập, hệ thống tự dẫn về /cms/pages (không phải /dashboard).')

H(2, 'Checklist tổng — Design phải biết làm')
check('Đăng nhập / đăng xuất')
check('Tạo và sửa Pages (trang landing page)')
check('Cập nhật ảnh và nội dung 4 Chương trình học')
check('Sửa Header Menu (label, link, thứ tự)')
check('Viết và đăng bài Blog/Tin tức')
check('Cập nhật Footer (hotline, mạng xã hội)')
check('Upload và quản lý Media Library')

H(2, '7.1. Sau đăng nhập — Landing page Design')
img('design_landing_cms_pages', 11, 'Sau login, Design vào thẳng Pages CMS')
P('Sidebar Design chỉ có 6 mục:')
bullet('Website CMS — quản lý các trang public')
bullet('Chương trình học — sửa 4 chương trình')
bullet('Header Menu — sửa menu chính')
bullet('Blog / Tin tức — viết và đăng bài')
bullet('Footer — sửa footer')
bullet('Media Library — upload và quản lý ảnh')

H(2, '7.2. Quy trình làm việc chuẩn')

H(3, 'Use case 1: Cập nhật trang chủ')
step(1, 'Vào Website CMS')
step(2, 'Click vào trang "Homepage"')
step(3, 'Trong Page Editor, sửa section hero, ảnh, text')
step(4, 'Lưu draft → preview')
step(5, 'Khi OK → Publish')

H(3, 'Use case 2: Đăng bài viết Tin tức')
img('design_cms_blog', 11, 'Blog admin (Design)')
step(1, 'Vào Blog / Tin tức')
step(2, 'Nhấn "Tạo bài mới"')
step(3, 'Nhập tiêu đề, slug, nội dung (rich text editor)')
step(4, 'Upload ảnh đại diện từ Media Library')
step(5, 'Đặt trạng thái Draft (lưu nháp) hoặc Published (đăng ngay)')
step(6, 'Lưu — bài published sẽ xuất hiện ở /tin-tuc và trang chủ')

H(3, 'Use case 3: Đổi ảnh đại diện chương trình học')
img('design_cms_programs', 11, 'Programs CMS')
step(1, 'Vào Media Library → Upload ảnh mới (kích thước khuyến nghị 1200×800)')
step(2, 'Copy URL ảnh')
step(3, 'Vào Chương trình học → click vào chương trình muốn sửa')
step(4, 'Dán URL ảnh vào field "Cover image"')
step(5, 'Lưu')

H(3, 'Use case 4: Sửa hotline footer')
img('design_cms_footer', 11, 'Footer CMS')
step(1, 'Vào Footer')
step(2, 'Sửa field "Hotline" hoặc "Email"')
step(3, 'Lưu — public website cập nhật ngay')

H(3, 'Use case 5: Thêm mục mới vào Header Menu')
img('design_cms_header_menu', 11, 'Header Menu CMS')
step(1, 'Vào Header Menu')
step(2, 'Nhấn "+" để thêm mục mới')
step(3, 'Nhập label (tên hiển thị) và URL')
step(4, 'Kéo thả để đổi vị trí trong menu')
step(5, 'Lưu')

H(2, '7.3. Media Library — Tools dùng chung')
img('design_media', 11, 'Media Library')
P('Cloudinary CDN — đảm bảo tải nhanh.')
bullet('Upload nhiều ảnh cùng lúc')
bullet('Tự động generate variants (thumbnail, full size)')
bullet('Copy URL để dùng trong CMS')
bullet('Recommend: nén ảnh dưới 500KB trước khi upload để tiết kiệm storage')

H(2, '7.4. Best practices cho Design')
bullet('Ảnh hero nên có 2 phiên bản: desktop (1920×900) và mobile (768×1024)')
bullet('Alt text BẮT BUỘC cho ảnh quan trọng — tốt cho SEO')
bullet('Đặt slug có dấu gạch ngang, không dấu (vd: "khoa-tieng-anh-phonics")')
bullet('Trước khi Publish, luôn test trên mobile (mở public URL bằng điện thoại)')
bullet('Bài blog mới nên có ít nhất 1 ảnh + 300 từ để chuẩn SEO')

page_break()

# ============== PHẦN 9. FAQ ==============
H(1, 'PHẦN 8. Câu hỏi thường gặp & Xử lý lỗi')

H(2, '8.1. Đăng nhập')
P('Q: Tôi không đăng nhập được, báo "Mật khẩu không đúng"', bold=True)
P('A: Kiểm tra Caps Lock. Nhấn "Quên mật khẩu" để reset qua email. Liên hệ Admin nếu vẫn không được.')

P('Q: Sau khi đăng nhập, tôi không thấy menu nào trong sidebar', bold=True)
P('A: Role của bạn chưa được active. Liên hệ Admin để kiểm tra trong Users & Roles.')

P('Q: Vì sao Design login lại vào /cms/pages thay vì /dashboard?', bold=True)
P('A: Theo phân quyền, Design không có quyền vào Dashboard CRM. Hệ thống tự chuyển về landing page phù hợp với role.')

H(2, '8.2. CRM')
P('Q: Sales tôi báo không thấy lead', bold=True)
P('A: Vào "Phân lead" với role Admin/Manager → kiểm tra lead đã được gán cho sales đó chưa. Sales chỉ thấy lead có assigned_to = self.')

P('Q: Một lead bị "Bị trả về" — làm gì?', bold=True)
P('A: Nghĩa là sales được phân không cập nhật trạng thái trong 24h. Reassign cho sales khác trong Phân lead.')

P('Q: Lead mới từ form public không thấy trong CRM', bold=True)
P('A: Refresh trang. Kiểm tra Firestore console. Kiểm tra form public có submit thành công (hiện message "METTA đã nhận thông tin").')

H(2, '8.3. CAPI / Tracking')
P('Q: Event không về Meta Events Manager', bold=True)
P('A: 1) Vào /capi → Test Event Panel → gửi test event. 2) Kiểm tra Test Event Code khớp. 3) Vào /capi/events xem log lỗi. 4) Verify Pixel ID và Access Token đúng.')

P('Q: Pixel ID hiện tại có vẻ là placeholder (123456789000000)', bold=True)
P('A: ĐÚNG — đó là giá trị placeholder. Trước khi chạy ads thật, cập nhật Pixel ID và Access Token thật từ Meta Events Manager.')

H(2, '8.4. CMS')
P('Q: Tôi sửa nội dung CMS nhưng public website không cập nhật', bold=True)
P('A: 1) Kiểm tra đã nhấn Save/Publish chưa. 2) Refresh public site (Ctrl+F5 để clear cache). 3) Kiểm tra trạng thái: phải là "published" mới hiển thị.')

P('Q: Upload ảnh báo lỗi', bold=True)
P('A: 1) File phải < 10MB. 2) Định dạng JPG/PNG/WEBP. 3) Cloudinary có thể đang giới hạn tài khoản free — liên hệ Admin.')

H(2, '8.5. Liên hệ hỗ trợ')
bullet('Lỗi nghiệp vụ: liên hệ Manager phụ trách')
bullet('Lỗi kỹ thuật / không truy cập được: liên hệ Admin nội bộ')
bullet('Tài khoản: liên hệ Admin để tạo/reset')

# ============== END ==============
P('')
P('—— HẾT TÀI LIỆU ——', italic=True)
P('METTA Academy User Manual v1.0 — 02/06/2026', italic=True)

out = r'D:\Web Metta\User_Manual_METTA_Academy_v1.0.docx'
doc.save(out)
print('Saved:', out)
